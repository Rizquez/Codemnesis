# MODULES (EXTERNAL)
# ---------------------------------------------------------------------------------------------------------------------
from __future__ import annotations
from docx import Document
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx.text.run import Run
# ---------------------------------------------------------------------------------------------------------------------

# MODULES (INTERNAL)
# ---------------------------------------------------------------------------------------------------------------------
# Get listed here!
# ---------------------------------------------------------------------------------------------------------------------

# OPERATIONS / CLASS CREATION / GENERAL FUNCTIONS
# ---------------------------------------------------------------------------------------------------------------------

def bold_core_module_names(path: str) -> None:
    """
    Applies bold formatting to the module name in *core modules* entries within an already generated 
    DOCX document, preserving the rest of the formatting.

    This function performs post-processing on the Word document to resolve a limitation of `docxtpl`: 
    it is not possible to apply partial styles (for example, bold only to part of the text) without modifying 
    the template.

    Args:
        path (str):
            Path to the DOCX file that must be modified in-place.
    """
    doc = Document(path)

    for paragraph in doc.paragraphs:
        text = paragraph.text or ''
        if ': referenced by' not in text:
            continue

        name, rest = text.split(':', 1)
        name = name.strip()
        rest = f':{rest}'  

        # The existing run normally has the correct format (size/font).
        base_run = paragraph.runs[0] if paragraph.runs else None

        # Save base format (if there are no runs, create a temporary one)
        if base_run is None:
            base_run = paragraph.add_run('')

        # Clear existing runs (without deleting the paragraph or its style)
        for r in paragraph.runs:
            r.text = ''

        # Create run for name (bold)
        run_bold = paragraph.add_run(name)
        _copy_run_format(base_run, run_bold)
        run_bold.bold = True

        # Create run for the rest (normal)
        run_normal = paragraph.add_run(rest)
        _copy_run_format(base_run, run_normal)
        run_normal.bold = base_run.bold # Normally None/False

    doc.save(path)

def _copy_run_format(src: Run, dest: Run) -> None:
    """
     Copies the typographic format from a source *run* to a destination.

    This method explicitly replicates character style attributes (font, size, color, and decorations) 
    to ensure that the reconstructed text maintains exactly the same visual appearance as the original 
    text in the document.

    It is used to avoid unintended changes in font size, typeface, or style when creating new runs with 
    `python-docx`.

    Args:
        src (Run):
            Reference run containing the original formatting.
        dest (Run):
            Destination run to which the formatting will be copied.
    """
    try:
        dest.style = src.style
    except Exception:
        pass # Einige Run-Stile können möglicherweise nicht direkt zugewiesen werden

    dest.font.name = src.font.name
    dest.font.size = src.font.size
    dest.font.color.rgb = src.font.color.rgb
    dest.font.highlight_color = src.font.highlight_color

    dest.font.italic = src.font.italic
    dest.font.underline = src.font.underline
    dest.font.all_caps = src.font.all_caps
    dest.font.small_caps = src.font.small_caps
    dest.font.shadow = src.font.shadow
    dest.font.outline = src.font.outline
    dest.font.strike = src.font.strike
    dest.font.subscript = src.font.subscript
    dest.font.superscript = src.font.superscript

# ---------------------------------------------------------------------------------------------------------------------
# END OF FILE